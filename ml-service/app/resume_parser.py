from io import BytesIO
from typing import Optional
import pdfplumber
import docx
import re


def _clean_text(text: str) -> str:
    """Clean text while preserving line structure"""
    if not text:
        return text
    
    # Collapse multiple spaces into single space
    text = re.sub(r' +', ' ', text)
    
    # Collapse more than 2 newlines into single newline
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract text from resume with improved layout-aware parsing"""
    name = filename.lower() if filename else ''
    text = ''
    try:
        if name.endswith('.pdf'):
            try:
                with pdfplumber.open(BytesIO(data)) as pdf:
                    # Extract from all pages
                    pages_text = []
                    tables_text = []
                    
                    for page in pdf.pages:
                        # Use layout-aware extraction to preserve structure
                        page_text = page.extract_text(x_tolerance=2, y_tolerance=2, layout=True)
                        if page_text:
                            pages_text.append(page_text)
                        
                        # Extract tables separately and append after normal text
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                table_text = ' '.join([' '.join([cell or '' for cell in row]) for row in table])
                                if table_text:
                                    tables_text.append(table_text)
                    
                    # Combine: normal text first, then tables
                    text = "\n".join(pages_text)
                    if tables_text:
                        text = text + "\n" + "\n".join(tables_text)
                    
                    # Clean up text while preserving structure
                    text = _clean_text(text)
                    
            except Exception as e:
                print(f"PDF extraction error: {e}")
                text = ''
        elif name.endswith('.docx'):
            try:
                bio = BytesIO(data)
                d = docx.Document(bio)
                
                all_text = []
                
                # Extract header text first (insert at TOP)
                for section in d.sections:
                    if section.header:
                        for p in section.header.paragraphs:
                            if p.text.strip():
                                all_text.append(p.text)
                
                # Extract from paragraphs
                for p in d.paragraphs:
                    if p.text.strip():
                        all_text.append(p.text)
                
                # Extract from tables and append after normal text
                table_text = []
                for table in d.tables:
                    for row in table.rows:
                        row_text = ' '.join([cell.text for cell in row.cells if cell.text.strip()])
                        if row_text:
                            table_text.append(row_text)
                
                text = "\n".join(all_text)
                if table_text:
                    text = text + "\n" + "\n".join(table_text)
                
                # Clean up text while preserving structure
                text = _clean_text(text)
                
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                text = ''
        elif name.endswith('.doc'):
            # Try to read as text (old .doc format - limited support)
            try:
                text = data.decode('utf-8', errors='ignore')
            except Exception:
                text = ''
        else:
            try:
                text = data.decode('utf-8', errors='ignore')
            except Exception:
                text = ''
    except Exception as e:
        print(f"Text extraction error: {e}")
        text = ''
    
    return text
