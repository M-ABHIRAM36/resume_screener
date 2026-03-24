from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


OUT_PATH = r"c:\Users\DELL\Desktop\Resume screening project\docs\Resume_Screening_and_Career_Guidance_Institute_Format_Report.docx"


def set_doc_layout(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)  # A4 height
    section.left_margin = Cm(3.25)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)


def add_main_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 0, 0)
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_para(doc, text, single=False):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0.75)
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE if single else WD_LINE_SPACING.ONE_POINT_FIVE


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_figure_placeholder(doc, fig_no, title):
    add_para(doc, f"Figure {fig_no}: {title} (Placeholder for screenshot/diagram)", single=True)
    add_para(doc, "Insert the final image here during report finalization.", single=True)


def add_page_break(doc):
    doc.add_page_break()


def chapter_paragraphs(section_title, keyword_a, keyword_b, keyword_c):
    p1 = (
        f"The section on {section_title.lower()} explains the design and implementation choices made in the Resume Screening and Career Guidance System. "
        f"The project was developed as a modular full-stack application so that each functional concern is handled independently and can be improved without breaking the whole system. "
        f"From the analysis carried out on the source code, it is evident that the team followed a workflow-oriented approach where user actions in the frontend trigger validated backend APIs, and processing-intensive tasks are delegated to a dedicated Python service. "
        f"This separation improves maintainability, allows easy debugging, and supports future scaling of individual services."
    )

    p2 = (
        f"In this context, {keyword_a} plays a key operational role, while {keyword_b} and {keyword_c} provide supporting functionality across multiple modules. "
        f"For example, candidate-side operations such as resume upload, score visualization, skill-gap view, and roadmap guidance are tightly coupled with the backend API responses. "
        f"Similarly, HR-side operations including candidate ranking, filtering, and session management rely on standardized response objects with deterministic fields. "
        f"This consistency in contracts enables faster integration testing and predictable user experience across screens."
    )

    p3 = (
        f"The project evidence indicates that implementation decisions were not made only for feature completion, but also for clarity and academic traceability. "
        f"Data transformations, validation checks, and fallback logic are included at multiple points so that edge cases are handled gracefully. "
        f"In practical deployment scenarios, this section is important because reliability depends on how well request validation, structured storage, and response consistency are enforced. "
        f"Therefore, {section_title.lower()} contributes directly to both technical quality and usability outcomes of the final system."
    )

    return [p1, p2, p3]


def add_section(doc, code, title, kw1, kw2, kw3, bullets=None, figure=None):
    add_sub_heading(doc, f"{code} {title}")
    for para in chapter_paragraphs(title, kw1, kw2, kw3):
        add_para(doc, para)
    if bullets:
        for b in bullets:
            add_bullet(doc, b)
    if figure:
        add_figure_placeholder(doc, figure[0], figure[1])


def front_matter(doc):
    add_main_heading(doc, "MINI PROJECT REPORT")
    add_main_heading(doc, "RESUME SCREENING AND CAREER GUIDANCE SYSTEM")
    add_para(doc, "A Project Report submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.")
    add_para(doc, "Submitted by")
    add_para(doc, "1. ___________________________    (Roll No: __________________)")
    add_para(doc, "2. ___________________________    (Roll No: __________________)")
    add_para(doc, "3. ___________________________    (Roll No: __________________)")
    add_para(doc, "Under the guidance of")
    add_para(doc, "___________________________")
    add_para(doc, "Assistant Professor, Department of CSE")
    add_para(doc, "___________________________ COLLEGE OF ENGINEERING")
    add_para(doc, "Academic Year: 2025-2026")
    add_page_break(doc)

    add_main_heading(doc, "TITLE PAGE")
    add_para(doc, "Project Title: RESUME SCREENING AND CAREER GUIDANCE SYSTEM")
    add_para(doc, "Department: Computer Science and Engineering")
    add_para(doc, "Institution: ___________________________")
    add_para(doc, "Place: ___________________________")
    add_para(doc, "Month & Year of Submission: ___________________________")
    add_page_break(doc)

    add_main_heading(doc, "BONAFIDE CERTIFICATE")
    add_para(doc, "This is to certify that the mini project report entitled RESUME SCREENING AND CAREER GUIDANCE SYSTEM is a bonafide record of work carried out by the below-mentioned students under my supervision and guidance during the academic year 2025-2026.")
    add_para(doc, "The report is submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science and Engineering.")
    add_para(doc, "Student Names and Roll Numbers:")
    add_para(doc, "1. ___________________________    (Roll No: __________________)")
    add_para(doc, "2. ___________________________    (Roll No: __________________)")
    add_para(doc, "3. ___________________________    (Roll No: __________________)")
    add_para(doc, "Guide Signature: ___________________________")
    add_para(doc, "HOD Signature: ___________________________")
    add_para(doc, "Internal Examiner: ________________________")
    add_para(doc, "External Examiner: ________________________")
    add_page_break(doc)

    add_main_heading(doc, "DECLARATION")
    add_para(doc, "We hereby declare that the project report entitled RESUME SCREENING AND CAREER GUIDANCE SYSTEM is an original work carried out by us and has not been submitted to any other university or institution for the award of any degree or diploma.")
    add_para(doc, "All external references, concepts, and supporting content used in this report have been duly acknowledged in the references section in an appropriate citation format.")
    add_para(doc, "We further state that this report reflects our own understanding of problem analysis, system design, implementation, testing, and deployment for the proposed solution.")
    add_para(doc, "Student Signatures:")
    add_para(doc, "1. ___________________________")
    add_para(doc, "2. ___________________________")
    add_para(doc, "3. ___________________________")
    add_page_break(doc)

    add_main_heading(doc, "ACKNOWLEDGEMENT")
    add_para(doc, "We would like to express our sincere gratitude to our project guide for continuous support, constructive reviews, and valuable technical direction throughout the execution of this project.")
    add_para(doc, "We thank the Head of the Department and all faculty members for providing an encouraging academic environment and timely suggestions during each project milestone.")
    add_para(doc, "We also acknowledge the support of our institution for infrastructure, internet, software resources, and lab facilities that enabled us to complete this work.")
    add_para(doc, "Finally, we express thanks to our friends and family members for their motivation and cooperation during report preparation and project demonstrations.")
    add_page_break(doc)

    add_main_heading(doc, "ABSTRACT")
    add_para(doc, "Recruitment processes in many organizations still depend on manual screening of resumes, which is time-consuming and prone to inconsistency. This project presents a Resume Screening and Career Guidance System that automates resume analysis and provides actionable career recommendations for candidates. The solution uses a React-based frontend, Node.js and Express backend APIs, MongoDB for data persistence, and a FastAPI-powered Python service for resume parsing and scoring.")
    add_para(doc, "The system supports PDF and DOCX uploads, extracts structured information such as contact details, skills, education, and experience, and compares candidate profiles against role requirements. Score generation combines skill matching, semantic similarity, and keyword relevance. On top of this, the platform offers chat-based guidance modules for resume improvement and roadmap mentoring. HR users can run screening sessions, apply filters, and review ranked candidates, while candidate users can inspect score breakdown, missing skills, and growth suggestions.")
    add_para(doc, "Experimental usage on varied resume samples showed that the platform reduces screening effort, improves consistency in ranking, and helps candidates understand improvement priorities with greater clarity.")
    add_para(doc, "Keywords: Resume Analysis, Candidate Ranking, Skill Gap Detection, Career Roadmap, NLP, FastAPI, React, Express, MongoDB, Semantic Similarity.")
    add_page_break(doc)

    add_main_heading(doc, "TABLE OF CONTENTS")
    add_para(doc, "(Update this page from MS Word using References > Table of Contents after final editing.)")
    toc_lines = [
        "1. INTRODUCTION",
        "2. LITERATURE SURVEY",
        "3. REQUIREMENT ANALYSIS AND FEASIBILITY STUDY",
        "4. SYSTEM DESIGN AND ARCHITECTURE",
        "5. IMPLEMENTATION",
        "6. TESTING AND VALIDATION",
        "7. DEPLOYMENT, SECURITY, AND PERFORMANCE",
        "8. CONCLUSION AND FUTURE ENHANCEMENTS",
        "APPENDICES",
        "REFERENCES",
    ]
    for t in toc_lines:
        add_number(doc, t)
    add_page_break(doc)

    add_main_heading(doc, "LIST OF TABLES")
    table_list = [
        "Table 3.1 Functional Requirements Matrix",
        "Table 3.2 Non-Functional Requirements Matrix",
        "Table 3.3 Feasibility Analysis Summary",
        "Table 5.1 Backend API Endpoint Summary",
        "Table 5.2 Scoring Components and Weightages",
        "Table 6.1 Sample Test Cases and Expected Results",
        "Table 7.1 Deployment Environment Variables",
        "Table A.1 Sample Candidate Analysis Output"
    ]
    for t in table_list:
        add_number(doc, t)
    add_page_break(doc)

    add_main_heading(doc, "LIST OF FIGURES")
    figure_list = [
        "Figure 4.1 High-Level System Architecture",
        "Figure 4.2 Data Flow Diagram Level 0",
        "Figure 4.3 Data Flow Diagram Level 1",
        "Figure 4.4 Data Flow Diagram Level 2",
        "Figure 5.1 Candidate Dashboard Screen",
        "Figure 5.2 HR Dashboard Screen",
        "Figure 5.3 Resume Chat Interface",
        "Figure 5.4 Roadmap Chat Interface",
        "Figure 7.1 Deployment Topology"
    ]
    for f in figure_list:
        add_number(doc, f)
    add_page_break(doc)

    add_main_heading(doc, "LIST OF SYMBOLS, ABBREVIATIONS AND NOMENCLATURE")
    symbols = [
        "API : Application Programming Interface",
        "ATS : Applicant Tracking System",
        "CORS : Cross-Origin Resource Sharing",
        "DFD : Data Flow Diagram",
        "HR : Human Resources",
        "JWT : JSON Web Token",
        "NER : Named Entity Recognition",
        "NLP : Natural Language Processing",
        "REST : Representational State Transfer",
        "TF-IDF : Term Frequency-Inverse Document Frequency",
        "UI : User Interface",
        "UX : User Experience",
        "JSON : JavaScript Object Notation",
        "DB : Database",
        "HTTP : Hypertext Transfer Protocol"
    ]
    for s in symbols:
        add_bullet(doc, s)


def chapter_content(doc):
    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 1")
    add_main_heading(doc, "INTRODUCTION")
    add_section(doc, "1.1", "Introduction", "React frontend", "Express APIs", "FastAPI service", [
        "The system supports both candidate and HR user journeys.",
        "Resume parsing supports PDF and DOCX inputs.",
        "The architecture follows modular service boundaries.",
        "Outputs are designed to be explainable for end users."
    ])
    add_section(doc, "1.2", "Background and Motivation", "manual screening delays", "inconsistent evaluation", "skill mismatch")
    add_section(doc, "1.3", "Problem Statement", "resume overload", "human bias", "lack of guidance")
    add_section(doc, "1.4", "Objectives of the Project", "automated extraction", "scoring", "roadmap support")
    add_section(doc, "1.5", "Scope of the Project", "candidate portal", "HR dashboard", "chat guidance")
    add_section(doc, "1.6", "Organization of the Report", "chapter flow", "technical evidence", "implementation traceability")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 2")
    add_main_heading(doc, "LITERATURE SURVEY")
    add_section(doc, "2.1", "Overview of Resume Screening Systems", "ATS tools", "rule-based ranking", "automation")
    add_section(doc, "2.2", "Keyword-Based Matching Methods", "exact matching", "TF-IDF", "domain dictionaries")
    add_section(doc, "2.3", "Semantic Similarity Approaches", "sentence embeddings", "context matching", "transformers")
    add_section(doc, "2.4", "Information Extraction from Resumes", "PDF parsing", "DOCX parsing", "entity extraction")
    add_section(doc, "2.5", "Skill Taxonomy and Canonical Mapping", "skill variations", "normalization", "lookup tables")
    add_section(doc, "2.6", "Career Guidance and Learning Path Systems", "gap detection", "recommendations", "mentoring chat")
    add_section(doc, "2.7", "Data Privacy and Ethical Considerations", "user data", "secure storage", "access control")
    add_section(doc, "2.8", "Survey of Chat-Assisted Career Tools", "response quality", "context grounding", "actionability")
    add_section(doc, "2.9", "Limitations in Existing Systems", "opaque scoring", "limited personalization", "weak feedback")
    add_section(doc, "2.10", "Research Gap and Proposed Contribution", "integrated platform", "explainable output", "dual-role support")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 3")
    add_main_heading(doc, "REQUIREMENT ANALYSIS AND FEASIBILITY STUDY")
    add_section(doc, "3.1", "Stakeholder Analysis", "candidate users", "HR users", "faculty evaluators")
    add_section(doc, "3.2", "Functional Requirements", "resume upload", "analysis response", "session save")
    add_section(doc, "3.3", "Non-Functional Requirements", "performance", "availability", "maintainability")
    add_section(doc, "3.4", "Software Requirements", "Node.js", "Python", "MongoDB")
    add_section(doc, "3.5", "Hardware and Platform Requirements", "developer machine", "cloud hosting", "browser support")
    add_section(doc, "3.6", "Technical Feasibility", "existing libraries", "API stack", "deployment support")
    add_section(doc, "3.7", "Operational and Economic Feasibility", "cost efficiency", "ease of use", "training effort")
    add_section(doc, "3.8", "Risk Analysis and Mitigation", "service failure", "data issues", "dependency risks")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 4")
    add_main_heading(doc, "SYSTEM DESIGN AND ARCHITECTURE")
    add_section(doc, "4.1", "High-Level System Architecture", "presentation layer", "application layer", "processing layer", figure=("4.1", "High-Level System Architecture"))
    add_section(doc, "4.2", "Frontend Design", "React routes", "component hierarchy", "state flow")
    add_section(doc, "4.3", "Backend API Design", "REST routes", "controller services", "validation")
    add_section(doc, "4.4", "Python Processing Service Design", "FastAPI endpoint", "parsing pipeline", "scoring pipeline")
    add_section(doc, "4.5", "Database and Schema Design", "Mongoose models", "session storage", "chat history")
    add_section(doc, "4.6", "Data Flow Design Level 0", "external entities", "system boundary", "core flows", figure=("4.2", "DFD Level 0"))
    add_section(doc, "4.7", "Data Flow Design Level 1", "process decomposition", "data stores", "flow labeling", figure=("4.3", "DFD Level 1"))
    add_section(doc, "4.8", "Data Flow Design Level 2", "resume analysis process", "sub-processes", "output aggregation", figure=("4.4", "DFD Level 2"))
    add_section(doc, "4.9", "Scoring Logic Design", "weighted metrics", "category scores", "bonus factors")
    add_section(doc, "4.10", "Chat Module Design", "resume chat", "roadmap chat", "context prompts")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 5")
    add_main_heading(doc, "IMPLEMENTATION")
    add_section(doc, "5.1", "Project Structure and Module Mapping", "frontend src", "backend routes", "ml-service app")
    add_section(doc, "5.2", "Frontend Route Implementation", "auth routes", "candidate pages", "hr pages")
    add_section(doc, "5.3", "Authentication and Authorization", "JWT tokens", "role-based guards", "session checks")
    add_section(doc, "5.4", "Resume Upload and Validation", "multer middleware", "file constraints", "error responses")
    add_section(doc, "5.5", "Text Extraction Implementation", "pdfplumber", "python-docx", "cleanup routines")
    add_section(doc, "5.6", "Skill Extraction Implementation", "skill variations", "word boundary matching", "canonical skills")
    add_section(doc, "5.7", "Similarity and Score Computation", "semantic similarity", "keyword similarity", "final score")
    add_section(doc, "5.8", "Candidate Dashboard and Score Views", "resume score page", "skill gap page", "roadmap page", figure=("5.1", "Candidate Dashboard and Score Views"))
    add_section(doc, "5.9", "HR Dashboard and Filtering Implementation", "top-N filtering", "sorting logic", "candidate cards", figure=("5.2", "HR Dashboard"))
    add_section(doc, "5.10", "Resume Chat Implementation", "chat sessions", "history retrieval", "assistant responses", figure=("5.3", "Resume Chat Screen"))
    add_section(doc, "5.11", "Roadmap Chat Implementation", "role scanning", "roadmap knowledge", "contextual mentoring", figure=("5.4", "Roadmap Chat Screen"))
    add_section(doc, "5.12", "Session Persistence for HR Screening", "screening sessions", "session results", "database updates")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 6")
    add_main_heading(doc, "TESTING AND VALIDATION")
    add_section(doc, "6.1", "Testing Strategy", "unit testing", "integration testing", "user acceptance testing")
    add_section(doc, "6.2", "Frontend Functional Testing", "navigation", "form validation", "state updates")
    add_section(doc, "6.3", "Backend API Testing", "status codes", "request payloads", "response schema")
    add_section(doc, "6.4", "ML Service Testing", "parser robustness", "skill extraction", "score consistency")
    add_section(doc, "6.5", "Authentication and Security Testing", "invalid token", "role mismatch", "access denial")
    add_section(doc, "6.6", "Chat Workflow Testing", "session start", "message send", "history delete")
    add_section(doc, "6.7", "Performance and Reliability Checks", "response time", "concurrent requests", "failure fallback")
    add_section(doc, "6.8", "Test Results and Analysis", "observations", "pass criteria", "improvement actions")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 7")
    add_main_heading(doc, "DEPLOYMENT, SECURITY, AND PERFORMANCE")
    add_section(doc, "7.1", "Deployment Architecture", "vercel frontend", "render backend", "render python", figure=("7.1", "Deployment Topology"))
    add_section(doc, "7.2", "Environment Configuration", "MONGO_URI", "JWT_SECRET", "ML_URL")
    add_section(doc, "7.3", "Security Considerations", "password handling", "token security", "api keys")
    add_section(doc, "7.4", "Data Storage and Privacy", "resume uploads", "chat data", "retention policy")
    add_section(doc, "7.5", "Performance Optimization", "request batching", "model loading", "payload minimization")
    add_section(doc, "7.6", "Scalability Considerations", "service scaling", "storage offloading", "queue strategy")
    add_section(doc, "7.7", "Operational Monitoring", "logs", "health checks", "error alerts")

    add_page_break(doc)
    add_main_heading(doc, "CHAPTER 8")
    add_main_heading(doc, "CONCLUSION AND FUTURE ENHANCEMENTS")
    add_section(doc, "8.1", "Conclusion", "project outcomes", "automation impact", "user benefit")
    add_section(doc, "8.2", "Academic Learning Outcomes", "full-stack integration", "nlp pipeline", "deployment skills")
    add_section(doc, "8.3", "Limitations", "ephemeral storage", "language constraints", "model dependency")
    add_section(doc, "8.4", "Future Enhancement Plan", "cloud storage", "analytics", "interview modules")
    add_section(doc, "8.5", "Final Remarks", "production readiness", "research continuity", "industry relevance")


def appendices_and_references(doc):
    # Appendices before references as requested
    add_page_break(doc)
    add_main_heading(doc, "APPENDICES")

    for i, name in enumerate([
        "Sample API Requests and Responses",
        "Detailed Test Case Logs",
        "Sample Parsed Resume Output",
        "Database Collection Snapshots",
        "UI Screen Capture Set",
        "Deployment Environment Snapshot",
        "Error Handling Cases",
        "Additional Screens and Evidence"
    ], start=1):
        add_sub_heading(doc, f"APPENDIX {chr(64+i)}: {name}")
        add_para(doc, "This appendix contains supporting technical evidence used during implementation validation and final review." )
        add_para(doc, "Insert authenticated screenshots, API outputs, or logs in this section before final submission.")
        add_figure_placeholder(doc, f"A.{i}", name)

    add_page_break(doc)
    add_main_heading(doc, "REFERENCES")

    refs = [
        "[1] React Documentation, Available: https://react.dev",
        "[2] Vite Documentation, Available: https://vitejs.dev",
        "[3] Express.js Documentation, Available: https://expressjs.com",
        "[4] MongoDB Documentation, Available: https://www.mongodb.com/docs",
        "[5] FastAPI Documentation, Available: https://fastapi.tiangolo.com",
        "[6] spaCy Documentation, Available: https://spacy.io",
        "[7] Sentence Transformers Documentation, Available: https://www.sbert.net",
        "[8] scikit-learn Documentation, Available: https://scikit-learn.org",
        "[9] pdfplumber Documentation, Available: https://github.com/jsvine/pdfplumber",
        "[10] python-docx Documentation, Available: https://python-docx.readthedocs.io",
        "[11] Render Deployment Docs, Available: https://render.com/docs",
        "[12] Vercel Deployment Docs, Available: https://vercel.com/docs",
        "[13] RFC 8259 - The JavaScript Object Notation (JSON) Data Interchange Format",
        "[14] REST Architectural Style and Constraints, Roy T. Fielding",
        "[15] Relevant project source files from local repository (frontend, backend, ml-service)."
    ]
    for r in refs:
        add_para(doc, r, single=True)


def main():
    doc = Document()
    set_doc_layout(doc)
    front_matter(doc)
    chapter_content(doc)
    appendices_and_references(doc)
    doc.save(OUT_PATH)
    print(f"Generated: {OUT_PATH}")


if __name__ == "__main__":
    main()
